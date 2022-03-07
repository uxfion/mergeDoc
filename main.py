import os
import re
from win32com import client as wc
from docx import Document
from docxcompose.composer import Composer

# ---------- 读取文件目录 -------------
dir = '2018级电气学院自动化专业工程技术实习'
cwd = os.getcwd()
file_dir_list = []
for root, dirs, files in os.walk(dir):
    for file in files:
        if file.endswith('.doc'):
            file_dir_list.append(os.path.join(cwd, root, file))
print(f'一共找到{len(file_dir_list)}个doc文件')
# for each in file_dir_list:
#     print(each)

# ------ 将doc另存为docx ----------
word = wc.Dispatch("Word.Application")
for file in file_dir_list:
    doc = word.Documents.Open(file)
    doc.SaveAs(f'{file}x', 12)
    doc.Close()
word.Quit()

# ---------- 读取docx文件路径 ----------
docx_list = []
for root, dirs, files in os.walk(dir):
    for file in files:
        if file.endswith('.docx'):
            docx_list.append(os.path.join(cwd, root, file))
print(f'一共找到{len(docx_list)}个docx文件')

# -------- 按照顺序排列，这步不同人的电脑上可能会有报错，需要自己调试 -------------
docx_list.sort(key=lambda l: int(re.search('.*?(\d+)\.\s.*?', l).group(1)))
# 使用正则表达式找到序号，自己根据自己文档名称找到排序方式
# tmp = re.search('.*?(\d+)\.\s.*?', docx_list[0]).group(1)
# print(tmp)
for each in docx_list:
    print(each)

# -------- 合并文档 ---------
# 填充分页符号文档
page_break_doc = Document()
page_break_doc.add_page_break()
# 定义新文档
target_doc = Document(docx_list[0])
target_composer = Composer(target_doc)
for i in range(len(docx_list)):
    # 跳过第一个作为模板的文件
    if i==0:
        continue
    # 填充分页符文档
    target_composer.append(page_break_doc)
    # 拼接文档内容
    f = docx_list[i]
    target_composer.append(Document(f))
# 保存目标文档
target_composer.save('all.docx')

